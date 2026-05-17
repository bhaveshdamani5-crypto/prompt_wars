"""
LexGuard AI — Document Parser & Validator Utility
===================================================
Handles file validation (type, size), raw text extraction via pdfplumber,
and prepares byte content for Gemini API processing.

Supported file types: PDF, PNG, JPG, JPEG, TXT
Max file size: 20 MB
"""

import os
import io
import uuid
import logging
from typing import Tuple, Optional

import pdfplumber
import docx

logger = logging.getLogger("LexGuard-AI.Parser")

# ──────────────────────────────────────────────
# Constants
# ──────────────────────────────────────────────

ALLOWED_EXTENSIONS: set[str] = {".pdf", ".docx", ".png", ".jpg", ".jpeg", ".txt"}
ALLOWED_MIME_TYPES: set[str] = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "image/png",
    "image/jpeg",
    "image/jpg",
    "text/plain",
}
MAX_FILE_SIZE_BYTES: int = 20 * 1024 * 1024  # 20 MB


# ──────────────────────────────────────────────
# Validation Functions
# ──────────────────────────────────────────────

def validate_file_extension(filename: str) -> str:
    """
    Validate that the file extension is in the allowed set.
    Returns the lowercase extension string.
    Raises ValueError if invalid.
    """
    ext: str = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Allowed types: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )
    logger.info(f"File extension validated: {ext}")
    return ext


def validate_file_size(file_bytes: bytes, filename: str) -> float:
    """
    Validate that the file does not exceed the maximum allowed size.
    Returns the file size in kilobytes.
    Raises ValueError if too large.
    """
    size_bytes: int = len(file_bytes)
    size_kb: float = round(size_bytes / 1024, 2)

    if size_bytes > MAX_FILE_SIZE_BYTES:
        max_mb: float = MAX_FILE_SIZE_BYTES / (1024 * 1024)
        raise ValueError(
            f"File '{filename}' is {size_kb:.1f} KB which exceeds the "
            f"maximum allowed size of {max_mb:.0f} MB."
        )
    logger.info(f"File size validated: {size_kb} KB for '{filename}'")
    return size_kb


def validate_content_type(content_type: Optional[str], filename: str) -> None:
    """
    Cross-check the upload Content-Type header against allowed MIME types.
    Logs a warning but does NOT block (extensions are the primary gate).
    """
    if content_type and content_type not in ALLOWED_MIME_TYPES:
        logger.warning(
            f"Content-Type '{content_type}' for '{filename}' is not in the "
            f"standard allowed list. Proceeding based on file extension."
        )


# ──────────────────────────────────────────────
# Unique Filename Generator
# ──────────────────────────────────────────────

def generate_safe_filename(original_filename: str) -> str:
    """
    Generate a unique, collision-free filename using UUID prefix.
    Preserves the original extension for downstream processing.
    """
    ext: str = os.path.splitext(original_filename)[1].lower()
    safe_name: str = f"{uuid.uuid4().hex[:12]}_{original_filename.replace(' ', '_')}"
    logger.info(f"Generated safe filename: {safe_name}")
    return safe_name


# ──────────────────────────────────────────────
# Raw Text Extraction (pdfplumber for PDFs)
# ──────────────────────────────────────────────

def extract_raw_text_from_pdf(file_bytes: bytes) -> Tuple[str, int]:
    """
    Extract raw text from a PDF using pdfplumber.
    Returns a tuple of (extracted_text, page_count).
    This is the fallback/preliminary extraction before Gemini simplification.
    """
    text_parts: list[str] = []
    page_count: int = 0

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            logger.info(f"PDF opened successfully: {page_count} pages detected.")

            for i, page in enumerate(pdf.pages):
                page_text: Optional[str] = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())
                    logger.debug(f"Page {i + 1}: extracted {len(page_text)} characters.")
                else:
                    # Page might be a scanned image — flag for Gemini vision
                    text_parts.append(f"[Page {i + 1}: No extractable text — scanned/image content detected]")
                    logger.warning(f"Page {i + 1}: No text layer found. Flagged for Gemini vision processing.")

                # Also extract text from tables on the page
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        for row in table:
                            if row:
                                row_text: str = " | ".join([cell or "" for cell in row])
                                if row_text.strip():
                                    text_parts.append(f"[TABLE ROW] {row_text}")

    except Exception as e:
        logger.error(f"pdfplumber extraction failed: {str(e)}")
        raise ValueError(f"Failed to read PDF content: {str(e)}")

    raw_text: str = "\n\n".join(text_parts)

    if not raw_text.strip() or all("[No extractable text" in p for p in text_parts):
        logger.warning("PDF appears to be entirely scanned/image-based. Full Gemini vision required.")

    return raw_text, page_count


def extract_raw_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file using python-docx.
    """
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs_text = [p.text for p in doc.paragraphs]
        
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text for cell in row.cells]
                paragraphs_text.append(" | ".join(row_cells))
                
        return "\n\n".join(paragraphs_text)
    except Exception as e:
        logger.error(f"python-docx extraction failed: {str(e)}")
        raise ValueError(f"Failed to read DOCX content: {str(e)}")


def extract_raw_text_from_txt(file_bytes: bytes) -> str:
    """
    Extract text from a plain text file with multi-encoding fallback.
    """
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        logger.warning("UTF-8 decode failed, falling back to latin-1 encoding.")
        return file_bytes.decode("latin-1")


def extract_raw_text(file_bytes: bytes, filename: str) -> Tuple[str, Optional[int]]:
    """
    Master extraction dispatcher.
    Routes to the correct extractor based on file extension.
    Returns (raw_text, page_count_or_none).

    For images (PNG/JPG/JPEG), returns empty string — these go directly
    to Gemini vision for intelligent extraction.
    """
    ext: str = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return extract_raw_text_from_pdf(file_bytes)
    elif ext == ".docx":
        text: str = extract_raw_text_from_docx(file_bytes)
        return text, None
    elif ext == ".txt":
        text: str = extract_raw_text_from_txt(file_bytes)
        return text, None
    elif ext in {".png", ".jpg", ".jpeg"}:
        # Images cannot have local text extraction — Gemini handles them directly
        logger.info(f"Image file detected ({ext}). Text extraction delegated to Gemini Vision.")
        return "", None
    else:
        raise ValueError(f"No extractor available for file type '{ext}'.")

